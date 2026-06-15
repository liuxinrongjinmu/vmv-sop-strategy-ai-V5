from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from io import BytesIO
from typing import Tuple
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import os
import re
import urllib.request
import tempfile
import logging

logger = logging.getLogger(__name__)

class ReportExportService:
    """
    报告导出服务
    支持Markdown、PDF、Word格式导出
    """
    
    def __init__(self):
        self.chinese_font = None
        self.font_initialized = False
    
    def _init_fonts(self):
        """初始化中文字体，支持 Windows 和 Linux (Railway)"""
        if self.font_initialized:
            return
        
        self.font_initialized = True
        self.chinese_font = None
        
        font_paths = [
            "C:/Windows/Fonts/msyh.ttc",
            "C:/Windows/Fonts/simhei.ttf",
            "C:/Windows/Fonts/simsun.ttc",
            "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
            "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf",
        ]
        
        for font_path in font_paths:
            if os.path.exists(font_path):
                try:
                    pdfmetrics.registerFont(TTFont('ChineseFont', font_path))
                    self.chinese_font = 'ChineseFont'
                    logger.info(f"字体初始化成功: {font_path}")
                    return
                except Exception as e:
                    logger.warning(f"字体注册失败 {font_path}: {e}")
                    continue
        
        logger.info("未找到系统中文字体，尝试下载 Noto Sans SC...")
        
        # 多个字体下载源，提高可用性
        font_urls = [
            "https://github.com/googlefonts/noto-cjk/raw/main/Sans/OTF/SimplifiedChinese/NotoSansSC-Regular.otf",
            "https://cdn.jsdelivr.net/gh/notofonts/noto-cjk@main/Sans/OTF/SimplifiedChinese/NotoSansSC-Regular.otf",
        ]
        
        try:
            font_dir = os.path.join(tempfile.gettempdir(), "fonts")
            os.makedirs(font_dir, exist_ok=True)
            font_path = os.path.join(font_dir, "NotoSansSC-Regular.otf")
            
            if not os.path.exists(font_path) or os.path.getsize(font_path) < 1000:
                for font_url in font_urls:
                    try:
                        logger.info(f"尝试从 {font_url} 下载字体...")
                        urllib.request.urlretrieve(font_url, font_path)
                        if os.path.exists(font_path) and os.path.getsize(font_path) > 1000:
                            logger.info("字体下载成功")
                            break
                    except Exception as e:
                        logger.warning(f"下载源失败: {e}")
                        continue
            
            if os.path.exists(font_path) and os.path.getsize(font_path) > 1000:
                pdfmetrics.registerFont(TTFont('ChineseFont', font_path))
                self.chinese_font = 'ChineseFont'
                logger.info("字体注册成功: NotoSansSC-Regular.otf")
            else:
                logger.warning("字体文件无效")
        except Exception as e:
            logger.error(f"字体下载失败: {e}")
            logger.warning("将使用默认字体 (可能不支持中文)")
        
        if not self.chinese_font:
            logger.warning("未找到中文字体，PDF 导出可能会出现乱码")
    
    def export_markdown(self, content: str, title: str) -> Tuple[bytes, str]:
        """
        导出Markdown格式
        """
        filename = f"{title}.md"
        return content.encode('utf-8'), filename
    
    def export_pdf(self, content: str, title: str) -> Tuple[bytes, str]:
        """
        导出PDF格式
        使用reportlab生成PDF
        """
        # 延迟初始化字体，避免服务启动时阻塞
        self._init_fonts()
        
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=2*cm,
            leftMargin=2*cm,
            topMargin=2*cm,
            bottomMargin=2*cm
        )
        
        styles = getSampleStyleSheet()
        
        font_name = self.chinese_font if self.chinese_font else 'Helvetica'
        
        title_style = ParagraphStyle(
            'ChineseTitle',
            parent=styles['Title'],
            fontName=font_name,
            fontSize=18,
            leading=24,
            alignment=TA_CENTER,
            spaceAfter=20
        )
        
        heading1_style = ParagraphStyle(
            'ChineseHeading1',
            parent=styles['Heading1'],
            fontName=font_name,
            fontSize=16,
            leading=20,
            spaceBefore=16,
            spaceAfter=10
        )
        
        heading2_style = ParagraphStyle(
            'ChineseHeading2',
            parent=styles['Heading2'],
            fontName=font_name,
            fontSize=14,
            leading=18,
            spaceBefore=12,
            spaceAfter=8
        )
        
        heading3_style = ParagraphStyle(
            'ChineseHeading3',
            parent=styles['Heading3'],
            fontName=font_name,
            fontSize=12,
            leading=16,
            spaceBefore=10,
            spaceAfter=6
        )
        
        body_style = ParagraphStyle(
            'ChineseBody',
            parent=styles['Normal'],
            fontName=font_name,
            fontSize=11,
            leading=16,
            alignment=TA_JUSTIFY,
            spaceBefore=4,
            spaceAfter=4
        )
        
        bullet_style = ParagraphStyle(
            'ChineseBullet',
            parent=body_style,
            leftIndent=20,
            bulletIndent=10
        )
        
        story = []
        
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 0.5*cm))
        
        lines = content.split('\n')
        
        table_lines = []
        in_table = False
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # 表格检测
            if line.startswith('|') and line.endswith('|'):
                if not in_table:
                    in_table = True
                    table_lines = []
                # 跳过分隔行 (|---|---|)
                if re.match(r'^\|[\s\-:|]+\|$', line):
                    continue
                table_lines.append(line)
                continue
            else:
                # 如果之前在表格中，现在退出了，渲染表格
                if in_table and table_lines:
                    self._render_pdf_table(table_lines, story, body_style, font_name)
                    table_lines = []
                    in_table = False
            
            line = self._escape_html(line)
            
            if line.startswith('# '):
                text = line[2:]
                story.append(Paragraph(text, heading1_style))
            elif line.startswith('## '):
                text = line[3:]
                story.append(Paragraph(text, heading2_style))
            elif line.startswith('### '):
                text = line[4:]
                story.append(Paragraph(text, heading3_style))
            elif line.startswith('- '):
                text = '• ' + line[2:]
                story.append(Paragraph(text, bullet_style))
            elif line.startswith('**') and line.endswith('**'):
                text = '<b>' + line.strip('*') + '</b>'
                story.append(Paragraph(text, body_style))
            else:
                story.append(Paragraph(line, body_style))
        
        # 循环结束后，如果还在表格中
        if in_table and table_lines:
            self._render_pdf_table(table_lines, story, body_style, font_name)
        
        doc.build(story)
        buffer.seek(0)
        
        filename = f"{title}.pdf"
        return buffer.read(), filename
    
    def _render_pdf_table(self, table_lines: list, story: list, body_style, font_name: str):
        """渲染Markdown表格为PDF Table"""
        if not table_lines:
            return
        
        # 解析表格数据
        rows = []
        for line in table_lines:
            cells = [cell.strip() for cell in line.strip('|').split('|')]
            rows.append(cells)
        
        if not rows:
            return
        
        # 确定列数
        num_cols = max(len(row) for row in rows)
        
        # 补齐列数
        for row in rows:
            while len(row) < num_cols:
                row.append('')
        
        # 转义HTML
        for i, row in enumerate(rows):
            for j, cell in enumerate(row):
                rows[i][j] = self._escape_html(cell)
        
        # 创建表格
        col_width = (A4[0] - 4*cm) / num_cols
        table = Table(rows, colWidths=[col_width]*num_cols)
        
        # 表格样式
        style_commands = [
            ('FONTNAME', (0, 0), (-1, -1), font_name),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('BACKGROUND', (0, 0), (-1, 0), colors.Color(0.9, 0.95, 1.0)),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.Color(0, 51, 102)),
            ('TOPPADDING', (0, 0), (-1, -1), 4),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ('LEFTPADDING', (0, 0), (-1, -1), 6),
            ('RIGHTPADDING', (0, 0), (-1, -1), 6),
        ]
        
        table.setStyle(TableStyle(style_commands))
        story.append(table)
        story.append(Spacer(1, 0.3*cm))
    
    def _escape_html(self, text: str) -> str:
        """转义HTML特殊字符"""
        text = text.replace('&', '&amp;')
        text = text.replace('<', '&lt;')
        text = text.replace('>', '&gt;')
        return text
    
    def export_docx(self, content: str, title: str) -> Tuple[bytes, str]:
        """
        导出Word格式
        统一格式：标题、副标题、正文使用一致的样式
        """
        doc = Document()
        
        title_heading = doc.add_heading(title, 0)
        title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_heading_font(title_heading, 22, RGBColor(0, 51, 102))
        
        lines = content.split('\n')
        
        table_lines_docx = []
        in_table_docx = False
        
        for line in lines:
            original_line = line
            line = line.strip()
            if not line:
                continue
            
            # 表格检测
            if line.startswith('|') and line.endswith('|'):
                if not in_table_docx:
                    in_table_docx = True
                    table_lines_docx = []
                if re.match(r'^\|[\s\-:|]+\|$', line):
                    continue
                table_lines_docx.append(line)
                continue
            else:
                if in_table_docx and table_lines_docx:
                    self._render_docx_table(table_lines_docx, doc)
                    table_lines_docx = []
                    in_table_docx = False
            
            if line.startswith('# '):
                heading = doc.add_heading(line[2:], 1)
                self._set_heading_font(heading, 18, RGBColor(0, 51, 102))
            elif line.startswith('## '):
                heading = doc.add_heading(line[3:], 2)
                self._set_heading_font(heading, 16, RGBColor(0, 76, 153))
            elif line.startswith('### '):
                heading = doc.add_heading(line[4:], 3)
                self._set_heading_font(heading, 14, RGBColor(0, 102, 153))
            elif line.startswith('- '):
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(line[2:])
                self._set_run_font(run, 11)
            elif line.startswith('**') and line.endswith('**'):
                p = doc.add_paragraph()
                run = p.add_run(line.strip('*'))
                run.bold = True
                self._set_run_font(run, 11)
            else:
                p = doc.add_paragraph()
                run = p.add_run(line)
                self._set_run_font(run, 11)
        
        # 循环结束后，如果还在表格中
        if in_table_docx and table_lines_docx:
            self._render_docx_table(table_lines_docx, doc)
        
        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        
        filename = f"{title}.docx"
        return docx_buffer.read(), filename
    
    def _render_docx_table(self, table_lines: list, doc):
        """渲染Markdown表格为Word Table"""
        if not table_lines:
            return
        
        rows = []
        for line in table_lines:
            cells = [cell.strip() for cell in line.strip('|').split('|')]
            rows.append(cells)
        
        if not rows:
            return
        
        num_cols = max(len(row) for row in rows)
        for row in rows:
            while len(row) < num_cols:
                row.append('')
        
        # 创建Word表格
        table = doc.add_table(rows=len(rows), cols=num_cols, style='Light Grid Accent 1')
        
        for i, row in enumerate(rows):
            for j, cell_text in enumerate(row):
                cell = table.cell(i, j)
                cell.text = cell_text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        self._set_run_font(run, 9)
        
        doc.add_paragraph()  # 表格后空行
    
    def _get_chinese_font_name(self) -> str:
        """获取可用的中文字体名称，跨平台兼容"""
        import platform
        if platform.system() == 'Windows':
            return '微软雅黑'
        else:
            return 'Noto Sans CJK SC'

    def _set_heading_font(self, heading, size: int, color: RGBColor):
        """设置标题字体"""
        font_name = self._get_chinese_font_name()
        for run in heading.runs:
            run.font.size = Pt(size)
            run.font.color.rgb = color
            run.font.bold = True
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    
    def _set_run_font(self, run, size: int):
        """设置正文字体"""
        font_name = self._get_chinese_font_name()
        run.font.size = Pt(size)
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

report_export_service = ReportExportService()
